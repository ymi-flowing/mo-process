"""
MO Process - SNS Multi Family Management LLC / ResMan

Runs the full Move-Out / Final Account Statement workflow for a former resident.

FLOW (in order)
  1. Login to ResMan.
  2. Open Move Out Reconciliation from the resident's Leasing Workflow
     (also captures `oid` = BillingAccountId from the MOR link's data-href —
      needed later as the `Id` field for the ResMan Partner API upload).
  3. Fill Move-out rec. date (defaults to today).
  4. Add each charge (default category: "Cleaning/Damage Charges").
  5. Capture MOR totals — including `depositsAvailableTotal` (Totals row of
     the Available Deposit table, which sums multi-deposit residents).
  6. Approve the reconciliation.
  7. Read resident name / unit / email + forwarding address (via the Lease
     Edit dialog for stable structured fields; falls back to Unit address
     if forwarding is blank).
  8. Generate a filled Claim Form docx from Claim Form Example.Docx.
  9. Download the ResMan-generated Final Account Statement PDF.
 10. Merge Claim Form + FAS → `Move Out Docs - Unit <#>.pdf`
     (strips empty "Images for Charges" page from the FAS).
 11. Docupost sendletter — pushes merged PDF to public repo, POSTs to
     https://app.docupost.com/api/1.1/wf/sendletter. Gated on the merged
     PDF existing (NOT on the resident email). Returns letter_id + cost.
 12. Capture cert screenshot from Docupost dashboard (fresh browser
     context; polls up to 180s for USPS tracking #; screenshots the
     .bubble-element.Group containing Delivery Status + Recipient + Sender).
     Saves as `out/MO Docs - Mail Certification.png`.
 13. ResMan API POST /Documents — upload merged PDF to /Move-Out Docs
     (HARD-FAIL: email step needs the file in Documents to attach).
 14. ResMan API POST /Documents — upload cert PNG to /Move-Out Docs
     (SOFT-FAIL: letter is already mailed).
 15. Open resident email, apply ***MO Docs Email template, attach the
     merged PDF from ResMan's picker (expands /Move-Out Docs folder;
     falls back to "Add from Computer" if the picker doesn't surface
     the file after retries). Click Send.
 16. Verify via Communication Log (Kendo grid lazy-hydrates — waits 10s
     then retries once at 15s).

USAGE
    python run_mo_process.py --payload @payload.json
    python run_mo_process.py --payload -                # read JSON from stdin
    python run_mo_process.py --payload @payload.json --no-send
    python run_mo_process.py --payload @payload.json --headless

PAYLOAD (single object)
    {
      "leaseUrl": "https://sns.myresman.com/#/Residents/Detail/<leaseId>",
      "charges": [
        { "description": "Cleaning",        "amount": 150.00 },
        { "description": "Carpet Cleaning", "amount": 200.00 }
      ],
      "morDate": "8/3/2026",              // optional; defaults to today (M/D/YYYY)
      "email": {
        "enabled": true,                  // default true
        "from": "property",               // "property" or "assistant"; default property
        "template": "***MO Docs Email"    // default "***MO Docs Email"
      },
      "docupost": {                       // optional; skipped if enabled=false
        "enabled": true,
        "class":        "usps_first_class",
        "servicelevel": "certified",
        "color":        false,
        "doublesided":  false,
        "sender": {                       // optional; defaults to properties.json
          "name": "49th St Apartments", "address1": "8400 49th Street N",
          "city": "Pinellas Park", "state": "FL", "zip": "33781"
        }
      },
      "outputDir": "out"                  // optional; defaults to CWD/out
    }

ENV VARIABLES
    RESMAN_USER, RESMAN_PASS                   — ResMan web login (Playwright).
    RESMAN_API_KEY, RESMAN_PARTNER_ID,         — ResMan Partner API (upload docs).
      RESMAN_ACCOUNT_ID                          Basic auth + Account-Id header.
    DOCUPOST_TOKEN                             — Docupost sendletter API.
    DOCUPOST_WEB_USER, DOCUPOST_WEB_PASS       — Docupost dashboard (cert screenshot).
    GITHUB_TOKEN                               — pushing merged PDF to repo/examples/.
                                                 (auto-provided by GH Actions)
    All creds also readable from Cardentials.txt (local dev fallback). Env wins.

RESULT
    A single JSON object is ALWAYS printed to stdout at the end (both on success
    and on error). Downstream (n8n, GitHub Actions) can rely on this shape:

    {
      "status": "sent" | "sent_no_email" | "parked" | "error",
      "startedAt": "2026-08-03T22:25:41Z",
      "endedAt":   "2026-08-03T22:31:47Z",
      "durationSeconds": 366,
      "resident": {
        "name": "Beatriz Pestana Gonzalez", "unit": "1511",
        "property": "49th St Apartments",
        "leaseUrl": "https://sns.myresman.com/#/Residents/Detail/...",
        "email": "..."
      },
      "mor": {
        "date": "8/3/2026", "status": "Complete",
        "charges":  [ {"category":"...","description":"...","amount":200} ],
        "totals":   {
          "currentOpenBalanceTotal": "0.00",
          "finalMoveOutChargesCreditsTotal": "420.00",
          "balanceBeforeDepositsTotal": "420.00",
          "depositsAvailableTotal": "799.00",
          "availableDepositApplied": "420.00",
          "paymentCreditRefund": "0.00",
          "depositRefund": "379.00",
          "balanceOwed": "0.00"
        },
        "forwardingAddress": { "street":"...", "city":"...", "state":"FL", "zip":"..." },
        "forwardingSource":  "resident" | "unit"
      },
      "docs": {
        "claimForm":              "out/Claim Form - <resident>.docx",
        "fasPdf":                 "out/Final Account Statement 8-3-2026 - <resident>.pdf",
        "combinedPdf":            "out/Move Out Docs - Unit <#>.pdf",
        "combinedPdfDocumentId":  "<ResMan documentId>"    // from POST /Documents
      },
      "email": {
        "attempted": true, "sent": true, "to": "...",
        "from": "property", "template": "***MO Docs Email",
        "subject": "... - Move-Out Documents",
        "attachedByResMan": [ {"name":"...","checked":true} ],
        "commLogVerified": true, "commLogRow": "..."
      },
      "docupost": {
        "letterId": "...", "cost": 12.33,
        "class": "usps_first_class", "servicelevel": "certified",
        "pdfUrl": "https://raw.githubusercontent.com/...",
        "certification": {
          "uploaded": true,
          "path": "out/MO Docs - Mail Certification.png",
          "documentId": "<ResMan documentId>",
          "tracking": "92xxx...",
          "error": null
        }
      },
      "github": { "repo": "ymi-flowing/mo-process", "runUrl": "..." },
      "logs":   [ "hh:mm:ss ...", ... ],
      "error":  null   // populated on failure with { "type", "message", "traceback" }
    }
"""
import argparse
import base64
import json
import os
import re
import sys
import shutil
import time
import traceback
from datetime import datetime, timezone
from pathlib import Path
from urllib.parse import quote as urllib_quote

from playwright.sync_api import sync_playwright, Page, TimeoutError as PWTimeout


# -------- Credentials (env overrides; defaults match SNS_Assistant) ---------
USERNAME = os.environ.get("RESMAN_USER") or "SNS_Assistant"
PASSWORD = os.environ.get("RESMAN_PASS") or "SNSassistant123$"
LOGIN_URL = "https://sns.myresman.com/"

DEFAULT_CATEGORY = "Cleaning/Damage Charges"
DEFAULT_TEMPLATE = "***MO Docs Email"

HERE = Path(__file__).parent.resolve()
CLAIM_TEMPLATE = HERE / "Claim Form Example.Docx"
PROPERTIES_FILE = HERE / "properties.json"

# --- Docupost + GitHub Contents API defaults ---
DOCUPOST_URL         = "https://app.docupost.com/api/1.1/wf/sendletter"
DOCUPOST_LOGIN_URL   = "https://app.docupost.com/login"
DOCUPOST_LETTER_URL  = "https://app.docupost.com/letter/{letter_id}"
DOCUPOST_EXAMPLES    = "examples"     # path prefix in the repo where Move Out Docs PDFs land
DOCUPOST_DEFAULTS = {
    "class":        "usps_first_class",
    "servicelevel": "certified",
    "color":        False,
    "doublesided":  False,
}
CERT_FILENAME        = "MO Docs - Mail Certification.png"

# --- ResMan Partner API (used to upload docs into resident's Documents tab) ---
RESMAN_API_URL       = "https://partners-api.myresman.com"
RESMAN_DOCS_FOLDER   = "/Move-Out Docs"
DEFAULT_SENDER = {
    "name":     "49th St Apartments",
    "address1": "8400 49th Street N",
    "city":     "Pinellas Park",
    "state":    "FL",
    "zip":      "33781",
}


# --------------------------- Properties directory --------------------------

def _load_properties() -> dict:
    """Load properties.json (property name -> config). Empty dict if missing."""
    if not PROPERTIES_FILE.exists():
        return {}
    try:
        return json.loads(PROPERTIES_FILE.read_text(encoding="utf-8"))
    except Exception as e:
        print(f"WARN: could not parse {PROPERTIES_FILE.name}: {e}", file=sys.stderr)
        return {}


PROPERTIES: dict = _load_properties()


def resolve_property(payload: dict, signals: dict) -> tuple[str, dict, dict]:
    """Return (property_name, property_config, signals).

    Signals ({'proid': ..., 'nameMatches': [...]}) are captured elsewhere on
    the resident detail page BEFORE navigation clobbers the header + MOR
    anchor's data-href. Resolution order:
      1. payload['property'] explicit override
      2. proid match against PROPERTIES entries with a non-null proid
      3. body-text name match against PROPERTIES keys (unique match required)
    Raises RuntimeError if none of the above resolves.
    """
    log(f"Property signals: proid={signals.get('proid')} nameMatches={signals.get('nameMatches')}")

    override = (payload or {}).get("property")
    if override:
        cfg = PROPERTIES.get(override)
        if not cfg:
            raise RuntimeError(f"payload.property={override!r} not in properties.json")
        return override, cfg, signals

    proid = signals.get("proid")
    if proid:
        for name, cfg in PROPERTIES.items():
            if cfg.get("proid") and cfg["proid"].lower() == proid.lower():
                return name, cfg, signals

    matches = signals.get("nameMatches") or []
    if len(matches) == 1:
        name = matches[0]
        return name, PROPERTIES[name], signals
    if len(matches) > 1:
        raise RuntimeError(
            f"Ambiguous property: page text matched multiple DB entries: {matches}. "
            f"Add proid to properties.json or set payload.property."
        )

    raise RuntimeError(
        f"Could not resolve property. proid={proid} nameMatches={matches}. "
        f"Add an entry to properties.json (and set proid={proid!r} if you have it) "
        f"or pass payload.property."
    )


# ------------------------------ Utilities ----------------------------------

_LOGS: list[str] = []


def log(msg):
    line = f"{datetime.now().strftime('%H:%M:%S')} {msg}"
    _LOGS.append(line)
    print(line, flush=True, file=sys.stderr)


def now_iso() -> str:
    return datetime.now(timezone.utc).strftime("%Y-%m-%dT%H:%M:%SZ")


def load_payload(arg):
    if arg == "-":
        raw = sys.stdin.read()
    elif arg.startswith("@"):
        raw = Path(arg[1:]).read_text(encoding="utf-8")
    else:
        raw = arg
    return json.loads(raw)


def today_str():
    n = datetime.now()
    return f"{n.month}/{n.day}/{n.year}"


def money(x):
    """Format a Decimal/float as ResMan-style '4,738.00'."""
    return f"{float(x):,.2f}"


def parse_money(s):
    if s is None:
        return 0.0
    return float(str(s).replace(",", "").strip())


def safe_slug(name):
    return re.sub(r"[^\w\-. ]+", "_", name).strip()


# ------------------------------ ResMan login -------------------------------

def login(page: Page):
    page.goto(LOGIN_URL, wait_until="domcontentloaded")
    log(f"Login user: {USERNAME!r}")
    try:
        page.wait_for_selector('input[name="Username"]', timeout=15000)
    except PWTimeout:
        if "myresman.com" in page.url and "Account/Login" not in page.url:
            log("Already authenticated.")
            return
        raise
    page.fill('input[name="Username"]', USERNAME)
    page.fill('input[name="Password"]', PASSWORD)
    page.click('button:has-text("Sign in")')
    page.wait_for_function(
        r"""() => /^https:\/\/sns\.myresman\.com\/#\//.test(location.href)
                  && !/Account\/Login/i.test(location.href)""",
        timeout=45000,
    )


# ------------------------------ MOR steps ----------------------------------

def open_move_out_rec(page: Page, lease_url: str, known_property_names: list | None = None) -> dict:
    """Navigate to the resident detail and click the visible Move Out Rec. link.

    Also captures property signals (proid + name matches) from the resident
    detail page BEFORE clicking Move Out Rec., because the click navigates
    away and the property header + MOR anchor's data-href both disappear.
    Returns {'dataHref': ..., 'propertySignals': {'proid': ..., 'nameMatches': [...]}}.
    """
    page.goto(lease_url, wait_until="domcontentloaded")
    # Wait for the sidebar's Move Out Rec. anchor to render.
    page.wait_for_function(
        "() => !!document.querySelector('#MoveOutReconciliationLink')",
        timeout=30000,
    )
    log("Resident detail loaded; capturing property signals + clicking Move Out Rec.")

    signals = page.evaluate(
        r"""(known) => {
          const out = { proid: null, nameMatches: [] };
          const a = document.querySelector('#MoveOutReconciliationLink, #MoveOutReconciliationOpenLink');
          if (a) {
            const href = a.getAttribute('data-href') || a.getAttribute('href') || '';
            const m = href.match(/proid=([0-9a-fA-F-]{36})/);
            if (m) out.proid = m[1];
          }
          const text = document.body.innerText || '';
          for (const name of (known || [])) {
            if (name && text.includes(name)) out.nameMatches.push(name);
          }
          return out;
        }""",
        known_property_names or [],
    )

    info = page.evaluate(
        r"""() => {
          const a = document.querySelector('#MoveOutReconciliationLink');
          const href = a.getAttribute('data-href') || '';
          window.jQuery(a).trigger('click');
          // Also extract 'oid' (BillingAccountId) from the data-href query
          // string — ResMan Partner API's POST /Documents needs it as `Id`
          // when type=Lease.
          const oidMatch = href.match(/[?&]oid=([0-9a-fA-F-]{36})/);
          return {
            dataHref: href,
            oid: oidMatch ? oidMatch[1] : null,
          };
        }"""
    )
    info['propertySignals'] = signals
    # Wait for the MOR page to render its Move-out rec. date input.
    page.wait_for_function(
        "() => !!document.getElementById('MoveOutReconciliationDate')",
        timeout=30000,
    )
    return info


def fill_mor_date(page: Page, mor_date: str):
    log(f"Setting Move-out rec. date = {mor_date}")
    page.evaluate(
        r"""(d) => {
          const el = document.getElementById('MoveOutReconciliationDate');
          window.jQuery(el).val(d).trigger('change').trigger('blur');
        }""",
        mor_date,
    )


def add_charge(page: Page, description: str, amount: float, category: str = DEFAULT_CATEGORY):
    """Click Add Charge / Credit, pick Category, fill Description + Amount."""
    log(f"Adding charge: {description} = ${amount} ({category})")

    row_ids_before = page.evaluate(
        r"""() => Array.from(document.querySelectorAll('input[name="MoveOutCharges.index"]')).map(i => i.value)"""
    )

    page.locator('button:has-text("Add Charge / Credit")').click()

    # Wait for a new row to appear.
    page.wait_for_function(
        r"""(before) => {
          const now = Array.from(document.querySelectorAll('input[name="MoveOutCharges.index"]')).map(i => i.value);
          return now.length > before.length;
        }""",
        arg=row_ids_before,
        timeout=15000,
    )
    row_id = page.evaluate(
        r"""(before) => {
          const now = Array.from(document.querySelectorAll('input[name="MoveOutCharges.index"]')).map(i => i.value);
          return now.find(v => !before.includes(v));
        }""",
        row_ids_before,
    )

    # Open the row's Category dropdown (2nd button on the row).
    page.evaluate(
        r"""(rowId) => {
          const trs = Array.from(document.querySelectorAll('tr')).filter(tr => tr.innerHTML.includes(rowId));
          const btns = Array.from(trs[0].querySelectorAll('button'));
          btns[1].click();
        }""",
        row_id,
    )

    # Pick the visible Category menu item.
    page.locator(f'[role="menuitem"]:visible:has-text("{category}")').first.click()

    # Fill Description + Amount.
    page.evaluate(
        r"""([rowId, desc, amt]) => {
          const desc_el = document.querySelector(`input[name="MoveOutCharges[${rowId}].Description"]`);
          const amt_el  = document.querySelector(`input[name="MoveOutCharges[${rowId}].ChargeAmount"]`);
          window.jQuery(desc_el).val(desc).trigger('change');
          window.jQuery(amt_el).val(amt).trigger('change').trigger('blur');
        }""",
        [row_id, description, f"{float(amount):.2f}"],
    )
    return row_id


def capture_mor_totals(page: Page) -> dict:
    """Read totals + deposit info from the MOR page before clicking Approve.

    `depositsAvailableTotal` = the TOTAL deposits-on-hand for the resident
    (sum of every deposit line — some residents have multiple). Read from
    the Totals row of the deposits table (the one whose header is
    "Available Deposit"). This is what the Claim Form's "Amount of
    Security Deposit" line should reflect — NOT `availableDepositApplied`,
    which is only the portion of the deposit applied to offset charges.
    """
    return page.evaluate(
        r"""() => {
          const grab = (id) => document.getElementById(id)?.value || null;
          const cells = Array.from(document.querySelectorAll('td, th'));
          const total = (label) => {
            const cell = cells.find(c => c.textContent.trim() === label);
            if (!cell) return null;
            const vals = Array.from(cell.parentElement.querySelectorAll('td')).map(td => td.textContent.trim()).filter(Boolean);
            return vals[vals.length - 1];
          };
          const depositsAvailableTotal = (() => {
            const th = Array.from(document.querySelectorAll('th')).find(t => t.textContent.trim() === 'Available Deposit');
            if (!th) return null;
            const table = th.closest('table');
            if (!table) return null;
            const totalsRow = Array.from(table.querySelectorAll('tr')).find(tr => {
              const first = tr.querySelector('td, th');
              return first && first.textContent.trim() === 'Totals';
            });
            if (!totalsRow) return null;
            const tds = Array.from(totalsRow.querySelectorAll('td, th'));
            // Column index of "Available Deposit" in the header row.
            const headerRow = th.parentElement;
            const headers = Array.from(headerRow.querySelectorAll('th, td'));
            const col = headers.indexOf(th);
            if (col < 0 || col >= tds.length) return null;
            return tds[col].textContent.trim();
          })();
          return {
            currentOpenBalanceTotal:         total('Current Open Balance Total'),
            finalMoveOutChargesCreditsTotal: total('Final Move Out Charges / Credits Total'),
            balanceBeforeDepositsTotal:      total('Balance before Deposits Total'),
            depositsAvailableTotal:          depositsAvailableTotal,
            availableDepositApplied:         document.querySelector('input[name*="ApplyToBalanceAmount"]')?.value || null,
            paymentCreditRefund:             grab('PaymentRefundAmount'),
            depositRefund:                   grab('CalculatedDepositRefundAmount'),
            balanceOwed:                     grab('BalanceOwed'),
          };
        }"""
    )


def approve_mor(page: Page):
    log("Actions -> Approve")
    page.locator('#Actions').click()
    page.locator('#Approve').click()
    # Approve redirects to /#/Residents/RedirectToDetail?ulgid=... and eventually
    # to the resident detail page. Wait for the Leasing Workflow to show Complete.
    page.wait_for_function(
        r"""() => document.body.innerText.includes('Move Out Rec (Complete)')""",
        timeout=45000,
    )
    log("MOR approved.")


# --------------------------- Forwarding address ----------------------------

def get_forwarding_address(page: Page) -> dict | None:
    """Read forwarding address by opening the Lease Edit dialog.

    Previously we screen-scraped the read-only "Forwarding address" cell,
    which ResMan renders in at least two DOM shapes (sibling divs / plain
    text with <br>) that collapsed under textContent and defeated the CSZ
    regex. The Edit dialog exposes the raw structured fields on stable ids:

      ForwardingAddress_StreetAddress  (textarea)
      ForwardingAddress_City
      ForwardingAddress_State
      ForwardingAddress_Zip
      ForwardingAddress_Country

    Zero parsing needed. We open the Lease-widget's Edit button, read the
    values, then close the dialog with its X (never Save). Returns None if
    the dialog can't be opened or every field is blank (caller falls back
    to the unit address).
    """
    opened = page.evaluate(
        r"""() => {
          const label = Array.from(document.querySelectorAll('label'))
            .find(l => l.textContent.trim().startsWith('Forwarding address'));
          if (!label) return false;
          // Walk up until we find a container that also holds the Lease
          // widget's Edit button (class 'dialog-form-link').
          let node = label;
          while (node && !(node.querySelector && node.querySelector('button.dialog-form-link'))) {
            node = node.parentElement;
          }
          const btn = node && node.querySelector('button.dialog-form-link');
          if (!btn) return false;
          btn.click();
          return true;
        }"""
    )
    if not opened:
        log("Forwarding: could not find Lease Edit button.")
        return None
    try:
        # Wait for the ForwardingAddress fields to render inside the newly
        # opened Lease dialog.
        page.wait_for_selector('#ForwardingAddress_StreetAddress', timeout=15000)
        vals = page.evaluate(
            r"""() => {
              const g = id => (document.getElementById(id)?.value || '').trim();
              return {
                street:  g('ForwardingAddress_StreetAddress'),
                city:    g('ForwardingAddress_City'),
                state:   g('ForwardingAddress_State'),
                zip:     g('ForwardingAddress_Zip'),
                country: g('ForwardingAddress_Country'),
              };
            }"""
        )
    except PWTimeout:
        log("Forwarding: Lease Edit dialog opened but fields never rendered.")
        vals = None
    finally:
        # Always close via the dialog's X (never Save) so we don't mutate
        # ResMan state. Robust against multiple stacked dialogs — we always
        # close the most recently opened visible ui-dialog.
        page.evaluate(
            r"""() => {
              const dialogs = Array.from(document.querySelectorAll('.ui-dialog'))
                .filter(d => d.getBoundingClientRect().width > 0);
              const d = dialogs[dialogs.length - 1];
              d?.querySelector('.ui-dialog-titlebar-close')?.click();
            }"""
        )
    if not vals or not vals.get("street"):
        return None
    unit_no = None
    m = re.search(r"\b(?:unit|apt|apartment|#)\s*([A-Za-z0-9\-]+)", vals["street"], re.I)
    if m:
        unit_no = m.group(1)
    return {
        "street":  vals["street"],
        "unitNo":  unit_no,
        "city":    vals["city"]  or None,
        "state":   vals["state"] or None,
        "zip":     vals["zip"]   or None,
        "county":  None,
        "country": vals["country"] or None,
    }


def parse_address_lines(text: str) -> dict:
    """Parse a multi-line address into street/city/state/zip/county."""
    lines = [ln.strip() for ln in text.splitlines() if ln.strip()]
    if not lines:
        return {}
    street = lines[0]
    unit_no = None
    m = re.search(r"\b(?:unit|apt|apartment|#)\s*([A-Za-z0-9\-]+)", street, re.I)
    if m:
        unit_no = m.group(1)

    # Try to find "City, ST ZIP" line.
    city, state, zipcode = None, None, None
    csz_rx = re.compile(r"^(.+?),\s*([A-Z]{2})\s*(\d{5}(?:-\d{4})?)$")
    for ln in lines[1:]:
        m = csz_rx.match(ln)
        if m:
            city, state, zipcode = m.group(1).strip(), m.group(2), m.group(3)
            break

    county = None
    for ln in lines[1:]:
        if re.search(r"county", ln, re.I):
            county = ln
            break

    return {
        "street": street,
        "unitNo": unit_no,
        "city": city,
        "state": state,
        "zip": zipcode,
        "county": county,
    }


def get_unit_address_via_new_tab(context, page: Page) -> dict:
    """Follow the Unit link on the resident page in a new tab and read Address.

    The Unit page renders the address label as `Address*` (with a required
    asterisk from a hidden <span class="required">). We match on the LABEL
    for="Address" — that's the ResMan-stable selector — and read the sibling
    `.fv` div for the value.
    """
    unit_href = page.evaluate(
        r"""() => {
          const links = Array.from(document.querySelectorAll('a[href*="/Units/Detail/"]'));
          return links[0] ? links[0].getAttribute('href') : null;
        }"""
    )
    if not unit_href:
        return {}
    url = unit_href if unit_href.startswith("http") else f"https://sns.myresman.com/{unit_href}"
    unit_page = context.new_page()
    try:
        unit_page.goto(url, wait_until="domcontentloaded")
        # Wait for the Address label (for="Address") to appear.
        unit_page.wait_for_function(
            r"""() => !!document.querySelector('label[for="Address"]')""",
            timeout=30000,
        )
        # Give the field value a moment to hydrate.
        unit_page.wait_for_function(
            r"""() => {
              const lbl  = document.querySelector('label[for="Address"]');
              const cell = lbl?.closest('td');
              const fv   = cell?.querySelector('.fv');
              return fv && fv.textContent.trim().length > 0;
            }""",
            timeout=15000,
        )
        addr_text = unit_page.evaluate(
            r"""() => {
              const lbl  = document.querySelector('label[for="Address"]');
              const cell = lbl?.closest('td');
              const fv   = cell?.querySelector('.fv');
              if (!fv) return null;
              // Prefer the structured child divs (street / city+state+zip / country).
              const parts = Array.from(fv.querySelectorAll('div, span'))
                .map(el => el.textContent.trim())
                .filter(t => t.length && t !== 'United States');
              if (parts.length) return parts.join('\n');
              return fv.textContent.trim();
            }"""
        )
    finally:
        unit_page.close()
    return parse_address_lines(addr_text or "")


# ----------------------------- Claim form gen ------------------------------

def generate_claim_form(
    out_dir: Path,
    resident_name: str,
    date_str: str,
    forwarding: dict,
    totals: dict,
    charges: list,
    property_config: dict | None = None,
) -> Path:
    """Fill Claim Form Example.Docx with resident's data. Handles deposit vs no-deposit.

    When ``property_config`` is provided (from properties.json), the property's
    return address (para 14) and Property Manager email (para 30) are rewritten
    to match. Para 28 (Management Signature) is always set to "Property Management".
    """
    from docx import Document
    from docx.oxml.ns import qn

    if not CLAIM_TEMPLATE.exists():
        raise FileNotFoundError(f"Claim template missing: {CLAIM_TEMPLATE}")

    dst = out_dir / f"Claim Form - {safe_slug(resident_name)}.docx"
    shutil.copy2(CLAIM_TEMPLATE, dst)

    # Field derivation (see docstring for the claim-form field map):
    #   sec_deposit       = deposits on hand (sum of every deposit line)
    #   credit_over       = credit-from-overpayment (only when past balance is negative)
    #   sec_plus_cred     = deposit + credit (money we're holding for the resident)
    #   dep_applied       = amount actually pulled from deposit to zero the balance
    #   total_charges     = move-out charges + prior owed balance (positive part only)
    #   landlord→resident = refund back to resident (deposit refund + payment refund)
    #   resident→landlord = balance still owed after applying deposit
    sec_deposit = parse_money(
        totals.get("depositsAvailableTotal")
        or totals.get("availableDepositApplied")
        or "0"
    )
    open_bal   = parse_money(totals.get("currentOpenBalanceTotal") or "0")
    mo_charges = parse_money(totals.get("finalMoveOutChargesCreditsTotal") or "0")
    dep_applied = parse_money(totals.get("availableDepositApplied") or "0")
    credit_over = abs(open_bal) if open_bal < 0 else 0.0
    sec_plus_cred = sec_deposit + credit_over
    total_charges = mo_charges + max(0.0, open_bal)
    resident_to_landlord = parse_money(totals.get("balanceOwed") or "0")
    dep_refund = parse_money(totals.get("depositRefund") or "0")
    pay_refund = parse_money(totals.get("paymentCreditRefund") or "0")
    landlord_to_resident = 0.0
    if dep_refund + pay_refund > 0 and resident_to_landlord == 0:
        landlord_to_resident = dep_refund + pay_refund

    street = forwarding.get("street") or ""
    city   = forwarding.get("city")
    state  = forwarding.get("state")
    zipc   = forwarding.get("zip")
    citystatezip = ", ".join([p for p in [city, f"{state} {zipc}".strip() if state or zipc else None] if p])

    d = Document(dst)
    paras = d.paragraphs

    def force(p, text):
        for i, run in enumerate(p.runs):
            run.text = text if i == 0 else ""
        if not p.runs:
            p.add_run(text)

    def set_paragraph_text(p, text):
        # python-docx's `p.runs` does NOT enumerate text inside <w:hyperlink>
        # children — so `force()` will leave a mailto: hyperlink intact and
        # you get "<new email><old email>" doubled up. Strip everything under
        # the paragraph except <w:pPr>, then add a fresh run.
        pPr = p._element.find(qn('w:pPr'))
        for child in list(p._element):
            if child is not pPr:
                p._element.remove(child)
        p.add_run(text)

    force(paras[6],  f"Date: {date_str}")
    force(paras[8],  f"Resident(s) Name:    {resident_name}")
    force(paras[9],  f"Address: {street}")
    force(paras[10], citystatezip)
    force(paras[12], f"This is a notice of my intention to impose a claim for damages in the amount of: $ {money(dep_applied)}")
    force(paras[18], f"Amount of Security Deposit:\t \t\t$ {money(sec_deposit)}")
    force(paras[19], f"Credit from Overpayment:\t\t \t$ {money(credit_over)}")
    force(paras[20], f"Total Security Deposit and Credit:\t \t$ {money(sec_plus_cred)}")
    force(paras[21], f"Total Charges:                \t \t\t$ {money(total_charges)}")
    force(paras[23], f"Total Due:  Landlord to Resident:              \t$ {money(landlord_to_resident)}")
    force(paras[24], f"                     Resident to Landlord:           \t$ {money(resident_to_landlord)}")

    if property_config:
        prop_line = (
            f"{property_config.get('address1', '')}, "
            f"{property_config.get('city', '')}, "
            f"{property_config.get('state', '')} "
            f"{property_config.get('zip', '')}"
        ).strip()
        set_paragraph_text(paras[14], prop_line)
        set_paragraph_text(
            paras[30],
            f"If you wish to dispute or disagree with any charges, you must submit "
            f"your request via email to the Property Manager at: {property_config.get('email', '')}",
        )

    # Remove the "Management Signature" paragraph entirely — no signature line.
    p28 = paras[28]
    p28._element.getparent().remove(p28._element)

    d.save(dst)
    log(f"Wrote claim form: {dst}")
    return dst


# ----------------------------- Documents I/O -------------------------------

def download_fas_pdf(page: Page, out_dir: Path, resident_name: str, date_str: str) -> Path | None:
    """Find the auto-generated Final Account Statement <date>.pdf and save it locally."""
    # Expand Documents accordion.
    page.evaluate(
        r"""() => {
          const h = Array.from(document.querySelectorAll('h3')).find(x => x.textContent.trim().startsWith('Documents'));
          h?.scrollIntoView({block:'center'});
          h?.click();
        }"""
    )
    page.wait_for_timeout(1500)
    info = page.evaluate(
        r"""() => {
          const el = Array.from(document.querySelectorAll('.document-name'))
            .find(x => x.textContent.trim().toLowerCase().includes('final account statement'));
          if (!el) return null;
          const row = el.closest('.document-row-grid');
          const dl  = row?.querySelector('a[href*="/Documents/Download"]');
          return { name: el.textContent.trim(), href: dl?.getAttribute('href') };
        }"""
    )
    if not info or not info.get("href"):
        log("Final Account Statement PDF not found on Documents tab.")
        return None

    log(f"Downloading FAS PDF: {info['name']}")
    payload = page.evaluate(
        r"""async (url) => {
          const res = await fetch(url, { credentials: 'include' });
          const buf = await res.arrayBuffer();
          const bytes = new Uint8Array(buf);
          let bin = '';
          const CHUNK = 32768;
          for (let i = 0; i < bytes.length; i += CHUNK) bin += String.fromCharCode.apply(null, bytes.subarray(i, i + CHUNK));
          return { status: res.status, size: bytes.length, b64: btoa(bin) };
        }""",
        info["href"],
    )
    data = base64.b64decode(payload["b64"])
    dst = out_dir / f"Final Account Statement {date_str.replace('/', '-')} - {safe_slug(resident_name)}.pdf"
    dst.write_bytes(data)
    log(f"Saved FAS PDF: {dst} ({len(data)} bytes)")
    return dst


def upload_document(page: Page, file_path: Path):
    """Click Add under Documents, pick file, click OK. Works for any file type."""
    log(f"Uploading document via Documents > Add: {file_path.name}")
    with page.expect_file_chooser() as fc_info:
        page.locator('button.add-files').click()
        page.wait_for_timeout(800)  # let the dialog render
        page.locator('input[type="file"]').click()
    fc = fc_info.value
    fc.set_files(str(file_path))
    # Wait for the Name field to auto-populate then click the dialog's OK.
    page.wait_for_timeout(1000)
    page.evaluate(
        r"""() => {
          const btns = Array.from(document.querySelectorAll('button')).filter(b => b.textContent.trim() === 'OK' && b.getBoundingClientRect().width>0);
          btns[0]?.click();
        }"""
    )
    # Wait for the dialog to close and the doc to appear.
    page.wait_for_function(
        r"""(fname) => Array.from(document.querySelectorAll('.document-name')).some(el => el.textContent.trim() === fname)""",
        arg=file_path.name,
        timeout=30000,
    )
    log(f"Uploaded: {file_path.name}")


# ------------------------------- Send Email --------------------------------

def open_send_email_dialog(page: Page):
    log("Opening resident email dialog.")
    # ResMan's sticky footer ("dock" icons) and BalancesCell can intercept
    # clicks on the mailto link in a smaller headless viewport. Skip
    # Playwright's hit-testing and fire the click via JS instead.
    hit = page.evaluate(
        r"""() => {
          const a = document.querySelector('a[href^="mailto:"]');
          if (!a) return { ok: false, reason: 'no mailto link' };
          a.scrollIntoView({ block: 'center' });
          a.click();
          return { ok: true, href: a.getAttribute('href') };
        }"""
    )
    if not hit or not hit.get("ok"):
        raise RuntimeError(f"Could not open email dialog: {hit}")
    page.wait_for_function(
        r"""() => !!document.getElementById('FromObject') && !!document.getElementById('Add')""",
        timeout=15000,
    )


def set_from(page: Page, preference: str):
    """preference: 'property' or 'assistant'. Safe to call multiple times —
    ResMan's template + attachment flow can silently reset From to the
    default Person, so we re-apply after both steps."""
    log(f"Setting From = {preference}")
    page.evaluate(
        r"""(pref) => {
          const sel = document.getElementById('FromObject');
          if (!sel) return { err: 'no #FromObject' };
          const opts = Array.from(sel.options);
          const match = pref === 'property'
            ? opts.find(o => o.dataset.objectType === 'Property')
            : opts.find(o => o.dataset.objectType === 'Person');
          if (!match) return { err: 'no match', options: opts.map(o => o.text) };
          sel.value = match.value;
          const display = document.getElementById('FromObjectInput');
          if (display) display.value = match.text;
          if (window.jQuery) {
            window.jQuery(sel).trigger('change');
            if (display) window.jQuery(display).trigger('change').trigger('autocompletechange');
          }
          return { selected: match.text };
        }""",
        preference,
    )


def apply_template(page: Page, template_name: str):
    log(f"Applying template: {template_name}")
    page.locator('button:has-text("Template")').click()
    # Templates render as anchors in a dialog; pick the visible one.
    page.locator(f'a:has-text("{template_name}")').first.click()
    # Wait for Subject to populate.
    page.wait_for_function(
        r"""() => (document.querySelector('input[name="Subject"], input#Subject')?.value || '').length > 0""",
        timeout=15000,
    )


def _open_attachment_picker(page: Page):
    """Click Add → Add from ResMan in the email dialog. Waits for the picker
    to render at least one document row."""
    page.evaluate(
        r"""() => {
          document.getElementById('Add')?.click();
        }"""
    )
    page.wait_for_timeout(400)
    page.evaluate(
        r"""() => {
          document.getElementById('btnAddFromCloud')?.click();
        }"""
    )
    page.wait_for_function(
        r"""() => !!document.querySelector('.document-name, .doc-name')""",
        timeout=15000,
    )


def _cancel_attachment_picker(page: Page):
    """Close the ResMan attachments picker via ESC, then verify the outer
    email dialog is still open.

    The earlier implementation clicked the topmost visible Cancel button. If
    the picker had already auto-dismissed, that Cancel hit the OUTER email
    dialog's Cancel — closing the whole compose. Every subsequent step
    (set_from, click_send) then silently no-op'd via ``?.click()`` and the
    runner logged "Email sent" without sending. Verified live (Luis Garcie
    T135, 7/14/2026): first-attempt run's Comm Log had zero rows for the
    intended send. ESC closes only the topmost modal; the verify raises so
    a silent no-op path can never follow."""
    page.keyboard.press("Escape")
    page.wait_for_timeout(600)
    if not page.evaluate("() => !!document.getElementById('FromObject')"):
        raise RuntimeError("attach picker close accidentally closed the outer email dialog")


def _wait_picker_has_files(page: Page, filenames: list, timeout_ms: int):
    """Wait until every requested filename is rendered as a .document-name
    in the currently-open attachments picker. Raises PWTimeout on miss."""
    page.wait_for_function(
        r"""(names) => {
          const rendered = Array.from(document.querySelectorAll('.document-name, .doc-name'))
            .map(el => el.textContent.trim());
          return names.every(n => rendered.includes(n));
        }""",
        arg=filenames,
        timeout=timeout_ms,
    )


def _expand_folder_in_picker(page: Page, folder_display_name: str) -> bool:
    """Click a `.folder-name` in the ResMan attachments picker to expand it,
    revealing the file rows + their checkboxes. Idempotent. Returns True if
    a folder with the given display name was found and clicked."""
    return page.evaluate(
        r"""(want) => {
          const el = Array.from(document.querySelectorAll('.folder-name'))
            .find(x => x.textContent.trim() === want);
          if (!el) return false;
          el.scrollIntoView({ block: 'center' });
          el.click();
          return true;
        }""",
        folder_display_name,
    )


def _check_files_in_picker(page: Page, filenames: list) -> list:
    """Tick the checkbox next to each filename in the visible picker.
    Returns [{name, checked, missing?}, ...]."""
    return page.evaluate(
        r"""(names) => {
          const results = [];
          names.forEach(name => {
            const els = Array.from(document.querySelectorAll('.document-name, span, div'))
              .filter(el => el.textContent.trim() === name && el.children.length === 0);
            for (const el of els) {
              const row = el.closest('.document-row-grid');
              if (!row) continue;
              const cb = row.querySelector('input[type="checkbox"]');
              if (cb && cb.getBoundingClientRect().width > 0) {
                if (!cb.checked) cb.click();
                results.push({ name, checked: cb.checked });
                return;
              }
            }
            results.push({ name, checked: false, missing: true });
          });
          return results;
        }""",
        filenames,
    )


def _attach_from_computer(page: Page, file_paths: list) -> None:
    """Fallback attach path when the ResMan picker doesn't surface our file:
    click 'Add' → 'Add from computer', pick the local file(s) via the file
    chooser. Used only when API-uploaded files aren't discoverable in the
    picker after folder expand + retries. Raises if the dropdown menu item
    can't be found — caller decides how to react."""
    log(f"Attaching via 'Add from Computer' (fallback): {[p.name for p in file_paths]}")
    with page.expect_file_chooser() as fc_info:
        page.evaluate(
            r"""() => {
              document.getElementById('Add')?.click();
            }"""
        )
        page.wait_for_timeout(400)
        # 'Add from computer' menu item — the dropdown renders it as a
        # visible link/button next to 'btnAddFromCloud'.
        page.evaluate(
            r"""() => {
              const cands = Array.from(document.querySelectorAll('a, button'))
                .filter(el => /add\s*from\s*computer/i.test((el.textContent || '').trim())
                           && el.getBoundingClientRect().width > 0);
              if (!cands.length) throw new Error('no Add from computer menu item');
              cands[0].click();
            }"""
        )
    fc = fc_info.value
    fc.set_files([str(p) for p in file_paths])
    page.wait_for_timeout(1500)


def attach_from_resman(
    page: Page,
    filenames: list,
    fallback_local_paths: list | None = None,
) -> list:
    """Attach each requested filename to the current Send Email dialog by
    opening the 'Add from ResMan' picker. Since we now API-upload docs into
    a `/Move-Out Docs` subfolder, we always try to expand that folder first
    so the file checkboxes are reachable.

    If the picker still doesn't have our file(s) after retries AND
    `fallback_local_paths` is provided, we close the picker and use 'Add
    from Computer' to attach the local files directly. Only raises if both
    strategies fail — the caller gets a status:error only when nothing
    could be attached.
    """
    log(f"Attaching from ResMan: {filenames}")

    last_result: list = []
    max_attempts = 3
    for attempt in range(1, max_attempts + 1):
        _open_attachment_picker(page)
        # Expand our /Move-Out Docs folder so file rows + checkboxes become
        # reachable. The folder-name text is the folder segment without the
        # leading '/'. Best-effort — a False here means the folder wasn't
        # (yet) rendered; we still let the picker settle and try the file
        # lookup, which will just report missing and we'll retry / fallback.
        folder_display = RESMAN_DOCS_FOLDER.lstrip("/")
        if _expand_folder_in_picker(page, folder_display):
            page.wait_for_timeout(500)

        # Wait up to 6s (first try) or 12s (subsequent) for the file names
        # to actually be rendered anywhere in the picker DOM.
        wait_ms = 6000 if attempt == 1 else 12000
        try:
            _wait_picker_has_files(page, filenames, timeout_ms=wait_ms)
        except PWTimeout:
            log(f"Picker attempt {attempt}: file(s) not indexed yet.")
            last_result = [{"name": n, "checked": False, "missing": True} for n in filenames]
            _cancel_attachment_picker(page)
            if attempt < max_attempts:
                page.wait_for_timeout(3000)
                continue
            # Exhausted picker attempts — try the local-file fallback if given.
            if fallback_local_paths:
                log("Picker never surfaced file(s); falling back to Add-from-Computer.")
                _attach_from_computer(page, fallback_local_paths)
                return [{"name": p.name, "checked": True, "source": "computer"} for p in fallback_local_paths]
            raise RuntimeError(
                f"Attachment(s) not found in ResMan picker after {max_attempts} attempts: {filenames}"
            )

        last_result = _check_files_in_picker(page, filenames)
        log(f"Attachment check result: {last_result}")

        missing = [r["name"] for r in last_result if r.get("missing") or not r.get("checked")]
        if not missing:
            break

        log(f"Picker attempt {attempt}: still missing/unchecked = {missing}. Retrying.")
        _cancel_attachment_picker(page)
        if attempt >= max_attempts:
            if fallback_local_paths:
                log("Picker checkbox still not selectable; falling back to Add-from-Computer.")
                _attach_from_computer(page, fallback_local_paths)
                return [{"name": p.name, "checked": True, "source": "computer"} for p in fallback_local_paths]
            raise RuntimeError(
                f"Attachment(s) still not selectable after {max_attempts} attempts: {missing}"
            )
        page.wait_for_timeout(3000)

    # Click OK on the picker to commit the selection. Trust the retry loop's
    # confirmed `checked: True` — ResMan's jQuery model has the attachment
    # committed at that point.
    page.evaluate(
        r"""() => {
          const btns = Array.from(document.querySelectorAll('button')).filter(b => b.textContent.trim() === 'OK' && b.getBoundingClientRect().width>0);
          btns[0]?.click();
        }"""
    )
    page.wait_for_timeout(600)
    return last_result

    return last_result


def click_send(page: Page):
    """Click Send in the resident email dialog with hard preconditions:
    the dialog must still be open (``#FromObject`` present) AND a visible
    Send button must exist. Prevents the silent no-op mode where the outer
    dialog was already closed by an earlier bug and ``?.click()`` did nothing
    while ``wait_for_function(...FromObject gone)`` returned immediately."""
    log("Verifying email dialog open before Send.")
    state = page.evaluate(
        r"""() => {
          if (!document.getElementById('FromObject')) return { ok: false, reason: 'email dialog closed (#FromObject missing)' };
          const btns = Array.from(document.querySelectorAll('button')).filter(b => b.textContent.trim() === 'Send' && b.getBoundingClientRect().width > 0);
          if (btns.length === 0) return { ok: false, reason: 'no visible Send button' };
          btns[0].click();
          return { ok: true, sendButtons: btns.length };
        }"""
    )
    log(f"Send click state: {state}")
    if not state.get("ok"):
        raise RuntimeError(f"Send preconditions failed: {state}")
    # Wait for the Send Email dialog to close (FromObject gone from the DOM).
    page.wait_for_function(
        r"""() => !document.getElementById('FromObject')""",
        timeout=45000,
    )
    log("Email dialog closed after Send.")


# ---------------------- Communication Log verification --------------------

def verify_via_comm_log(page: Page, lease_url: str, subject_hint: str, wait_seconds: int = 10) -> dict:
    """After Send: open the resident detail's Communication Log accordion,
    wait long enough for its Kendo grid to lazy-hydrate, then look for a
    row whose text contains today's subject (property name is enough).

    Returns {"verified": bool, "row": <text|None>, "opened": bool}. Never
    raises — the caller decides how to react (we downgrade status but do
    not fail the run, since a slow grid can produce a false negative)."""
    log(f"Verifying send via Communication Log (waiting {wait_seconds}s for Kendo grid to hydrate).")
    out = {"verified": False, "row": None, "opened": False}
    try:
        page.goto(lease_url, wait_until="domcontentloaded")
        page.wait_for_function(
            r"""() => !!document.querySelector('a[href^="mailto:"]')""",
            timeout=30000,
        )
        opened = page.evaluate(
            r"""() => {
              const hdrs = Array.from(document.querySelectorAll('h3, .accordion-header, .k-header, button, a'));
              const hdr = hdrs.find(h => (h.textContent || '').trim().startsWith('Communication Log'));
              if (!hdr) return { ok: false, reason: 'no Communication Log header' };
              hdr.scrollIntoView({ block: 'center' });
              hdr.click();
              return { ok: true };
            }"""
        )
        out["opened"] = bool(opened.get("ok"))
        page.wait_for_timeout(wait_seconds * 1000)
        row = page.evaluate(
            r"""(hint) => {
              const trs = Array.from(document.querySelectorAll('tr'));
              for (const tr of trs) {
                const t = (tr.innerText || '').replace(/\s+/g, ' ').trim();
                if (t.includes(hint) && t.includes('Email')) return t;
              }
              return null;
            }""",
            subject_hint,
        )
        if row:
            out["verified"] = True
            out["row"] = row
    except Exception as e:
        log(f"verify_via_comm_log soft-fail: {type(e).__name__}: {e}")
    log(f"Comm Log verification: {out}")
    return out


# ------------------------------ Docupost ----------------------------------

def _b64_file(p: Path) -> str:
    return base64.b64encode(p.read_bytes()).decode("ascii")


def push_pdf_to_repo(pdf_path: Path, repo: str, token: str, target_dir: str = DOCUPOST_EXAMPLES) -> str:
    """Push a PDF to `<repo>/<target_dir>/<pdf_path.name>` via the GitHub
    Contents API and return the public raw.githubusercontent.com URL.

    Uses PUT with a base64-encoded body. Overwrites the file if it already
    exists (fetches its SHA first). The repo must be public for Docupost's
    fetcher to reach the raw URL.

    Raises RuntimeError on non-2xx.
    """
    try:
        import requests  # type: ignore
    except ImportError as e:
        raise RuntimeError(f"requests library missing (needed for GitHub push): {e}")

    filename = pdf_path.name
    path     = f"{target_dir}/{filename}"
    api      = f"https://api.github.com/repos/{repo}/contents/{urllib_quote(path)}"
    headers  = {
        "Accept":        "application/vnd.github+json",
        "Authorization": f"Bearer {token}",
        "User-Agent":    "MO-Process-Runner/1.0",
    }

    # If the file already exists, we need its SHA to overwrite.
    sha = None
    r_get = requests.get(api + "?ref=main", headers=headers, timeout=30)
    if r_get.status_code == 200:
        sha = r_get.json().get("sha")

    body = {
        "message": f"add(examples): {filename}",
        "content": _b64_file(pdf_path),
        "branch":  "main",
    }
    if sha:
        body["sha"] = sha

    r = requests.put(api, headers=headers, json=body, timeout=120)
    if r.status_code >= 300:
        raise RuntimeError(f"GitHub push failed HTTP {r.status_code}: {r.text[:300]}")

    raw_url = f"https://raw.githubusercontent.com/{repo}/main/{target_dir}/{urllib_quote(filename)}"
    log(f"Pushed to repo: {path}")
    return raw_url


def _wait_raw_url_live(url: str, tries: int = 5, sleep_s: float = 2.0) -> bool:
    try:
        import requests  # type: ignore
    except ImportError:
        return False
    for i in range(1, tries + 1):
        try:
            r = requests.head(url, allow_redirects=True, timeout=30)
            if r.status_code == 200:
                return True
        except Exception:
            pass
        log(f"raw URL not ready (attempt {i}/{tries}); waiting {sleep_s}s")
        time.sleep(sleep_s)
    return False


def send_via_docupost(cfg: dict, sender: dict, recipient: dict, pdf_url: str, token: str) -> dict:
    """POST the Docupost sendletter request. Returns {letterId, cost, ...}.
    Raises RuntimeError on API error."""
    try:
        import requests  # type: ignore
    except ImportError as e:
        raise RuntimeError(f"requests library missing (needed for Docupost): {e}")

    params = {
        "api_token":    token,
        "pdf":          pdf_url,
        "class":        cfg.get("class",        DOCUPOST_DEFAULTS["class"]),
        "servicelevel": cfg.get("servicelevel", DOCUPOST_DEFAULTS["servicelevel"]),
        "color":        str(bool(cfg.get("color",       DOCUPOST_DEFAULTS["color"]))).lower(),
        "doublesided":  str(bool(cfg.get("doublesided", DOCUPOST_DEFAULTS["doublesided"]))).lower(),
        "description":  cfg.get("description", "")[:40],
        "from_name":     sender["name"],
        "from_address1": sender["address1"],
        "from_city":     sender["city"],
        "from_state":    sender["state"],
        "from_zip":      sender["zip"],
        "to_name":       recipient["name"],
        "to_address1":   recipient["address1"],
        "to_city":       recipient["city"],
        "to_state":      recipient["state"],
        "to_zip":        recipient["zip"],
    }
    if recipient.get("address2"):
        params["to_address2"] = recipient["address2"]
    if sender.get("address2"):
        params["from_address2"] = sender["address2"]

    r = requests.post(
        DOCUPOST_URL, params=params,
        headers={"Accept": "application/json", "User-Agent": "MO-Process-Runner/1.0"},
        timeout=180,
    )
    if r.status_code >= 300:
        raise RuntimeError(f"Docupost HTTP {r.status_code}: {r.text[:400]}")

    try:
        obj = r.json()
    except Exception:
        raise RuntimeError(f"Docupost response not JSON: {r.text[:400]}")

    letter_id = obj.get("letter_id") or obj.get("letterId")
    cost      = obj.get("cost")
    if not letter_id:
        raise RuntimeError(f"Docupost response missing letter_id: {obj}")
    log(f"Docupost queued letter_id={letter_id} cost=${cost}")
    return {
        "letterId":     letter_id,
        "cost":         cost,
        "class":        params["class"],
        "servicelevel": params["servicelevel"],
        "pdfUrl":       pdf_url,
    }


def _load_docupost_web_creds() -> tuple[str | None, str | None]:
    """Resolve Docupost dashboard credentials. Env wins over Cardentials.txt
    so CI can override with GitHub Secrets; local dev falls back to the file."""
    user = os.environ.get("DOCUPOST_WEB_USER")
    pw   = os.environ.get("DOCUPOST_WEB_PASS")
    if user and pw:
        return user, pw
    cred_file = HERE / "Cardentials.txt"
    if not cred_file.exists():
        return user, pw
    try:
        text = cred_file.read_text(encoding="utf-8")
    except Exception:
        return user, pw
    m = re.search(r"Docupost Web Login.*?Username:\s*(\S+).*?Password:\s*(\S+)", text, re.S | re.I)
    if not m:
        return user, pw
    return user or m.group(1).strip(), pw or m.group(2).strip()


def _load_resman_api_creds() -> tuple[str | None, str | None, str | None]:
    """Resolve ResMan Partner API credentials. Env wins over Cardentials.txt.
    Returns (partner_id, api_key, account_id) — any may be None if unresolved."""
    pid = os.environ.get("RESMAN_PARTNER_ID")
    key = os.environ.get("RESMAN_API_KEY")
    acc = os.environ.get("RESMAN_ACCOUNT_ID")
    if pid and key and acc:
        return pid, key, acc
    cred_file = HERE / "Cardentials.txt"
    if not cred_file.exists():
        return pid, key, acc
    try:
        text = cred_file.read_text(encoding="utf-8")
    except Exception:
        return pid, key, acc
    section = re.search(r"ResMan API:.*?(?=\n\S|\Z)", text, re.S | re.I)
    if not section:
        return pid, key, acc
    body = section.group(0)
    m_pid = re.search(r"Partner ID:\s*(\S+)", body, re.I)
    m_key = re.search(r"API Key:\s*(\S+)",    body, re.I)
    m_acc = re.search(r"Account ID:\s*(\S+)", body, re.I)
    return (
        pid or (m_pid.group(1).strip() if m_pid else None),
        key or (m_key.group(1).strip() if m_key else None),
        acc or (m_acc.group(1).strip() if m_acc else None),
    )


def upload_document_via_api(
    file_path: Path,
    object_id: str,
    property_id: str,
    folder: str = RESMAN_DOCS_FOLDER,
    show_in_resident_portal: bool = False,
) -> dict:
    """Upload a file to a resident's Documents via ResMan Partner API.

    POST https://partners-api.myresman.com/Documents  (multipart/form-data)
    Body fields:
      propertyId (property GUID), Id (BillingAccountId/oid), type=Lease,
      fileName, file, path (folder like '/Move-Out Docs').

    Retries on Cloudflare 5xx (502/503/504) up to 3 times with backoff, since
    the origin occasionally returns a transient 502 with a `retry_after: 60`
    hint (observed 2026-08-04 during initial API smoke test).

    Raises RuntimeError on unrecoverable non-2xx; returns the response JSON
    (including `documentId`) on success.
    """
    try:
        import requests  # type: ignore
    except ImportError as e:
        raise RuntimeError(f"requests library missing (needed for ResMan API): {e}")

    partner_id, api_key, account_id = _load_resman_api_creds()
    if not (partner_id and api_key and account_id):
        raise RuntimeError("ResMan API creds missing (env RESMAN_PARTNER_ID/API_KEY/ACCOUNT_ID or Cardentials.txt).")

    url = f"{RESMAN_API_URL}/Documents"
    headers = {"ResMan-Account-Id": account_id, "Accept": "application/json"}
    data = {
        "propertyId":           property_id,
        "Id":                   object_id,
        "type":                 "Lease",
        "fileName":             file_path.name,
        "path":                 folder,
        "showInResidentPortal": "true" if show_in_resident_portal else "false",
    }

    RETRIABLE = {502, 503, 504}
    last_err = None
    for attempt in range(1, 4):
        with open(file_path, "rb") as fh:
            files = {"file": (file_path.name, fh.read(), "application/octet-stream")}
        try:
            r = requests.post(
                url, auth=(partner_id, api_key),
                headers=headers, data=data, files=files, timeout=120,
            )
        except requests.RequestException as e:
            last_err = f"network error: {e}"
            log(f"ResMan upload attempt {attempt} transport error: {e}")
            if attempt < 3:
                time.sleep(30)
                continue
            raise RuntimeError(f"ResMan upload failed (transport): {last_err}")

        if 200 <= r.status_code < 300:
            try:
                obj = r.json()
            except Exception:
                obj = {"raw": r.text[:400]}
            log(f"ResMan upload OK: {file_path.name} → docID {obj.get('documentId')} (path={folder})")
            return obj

        if r.status_code in RETRIABLE and attempt < 3:
            log(f"ResMan upload attempt {attempt} got HTTP {r.status_code}; retrying in 65s.")
            time.sleep(65)
            continue
        raise RuntimeError(f"ResMan upload failed HTTP {r.status_code}: {r.text[:400]}")

    raise RuntimeError(f"ResMan upload failed after retries: {last_err}")


def capture_docupost_certification(
    browser,
    letter_id: str,
    out_dir: Path,
    web_user: str,
    web_pass: str,
    poll_seconds: int = 180,
) -> tuple[Path | None, str | None]:
    """Log into the Docupost dashboard, open the letter, wait for the USPS
    tracking # to render, and screenshot the middle panel (delivery status +
    order summary + recipient + sender) to `out_dir/MO Docs - Mail
    Certification.png`. Returns (path, tracking_number) — either may be None
    on soft failure. NEVER raises; callers log and move on."""
    from playwright.sync_api import Error as PWError
    cert_path = out_dir / CERT_FILENAME
    context = browser.new_context(viewport={"width": 1400, "height": 1200})
    try:
        page = context.new_page()
        page.goto(DOCUPOST_LETTER_URL.format(letter_id=letter_id), wait_until="domcontentloaded")
        # Docupost either (a) 302s to /login, or (b) already-logged-in
        # session lands us straight on the letter page. Detect which by
        # racing the login-form Email field against the 'Delivery Status'
        # text — whichever appears first wins.
        try:
            page.wait_for_function(
                r"""() => {
                  const hasLogin = !!Array.from(document.querySelectorAll('input'))
                    .find(i => (i.placeholder || i.name || '').toLowerCase().includes('email'));
                  const hasLetter = (document.body.innerText || '').includes('Delivery Status');
                  return hasLogin || hasLetter;
                }""",
                timeout=20000,
            )
        except PWError as e:
            log(f"Docupost: neither login form nor letter page rendered: {e}")
        # If it's the login page, fill and submit, then wait for the
        # post-login redirect to complete BEFORE we start polling. Not
        # waiting is what caused the "Execution context was destroyed"
        # race in run 30277568782.
        try:
            email_box = page.get_by_role("textbox", name="Email")
            if email_box.count() > 0 and email_box.first.is_visible():
                email_box.first.fill(web_user, timeout=10000)
                page.get_by_role("textbox", name="Password").fill(web_pass)
                page.get_by_role("button", name="Log in").click()
                # Wait for the letter page to actually render post-login.
                # networkidle is best-effort; the Delivery Status text is the
                # authoritative signal that we're on the right page and JS
                # rendering has completed.
                try:
                    page.wait_for_load_state("networkidle", timeout=15000)
                except PWError:
                    pass
                page.wait_for_function(
                    r"""() => (document.body.innerText || '').includes('Delivery Status')""",
                    timeout=20000,
                )
        except PWError as e:
            log(f"Docupost login step issue (continuing): {e}")
        # Poll for the tracking #. Each iteration is wrapped so a transient
        # nav (Bubble.io does async rerenders) doesn't blow up the loop.
        tracking = None
        deadline = time.time() + poll_seconds
        while time.time() < deadline:
            try:
                tracking = page.evaluate(
                    r"""() => {
                      const m = (document.body.innerText || '').match(/USPS Tracking #\s*(\d{18,})/);
                      return m ? m[1] : null;
                    }"""
                )
            except PWError as e:
                log(f"Docupost poll iteration transient error (retrying): {e}")
                tracking = None
            if tracking:
                break
            page.wait_for_timeout(2000)
        if tracking:
            log(f"Docupost tracking #: {tracking}")
        else:
            log(f"Docupost tracking # not visible after {poll_seconds}s — screenshotting current state.")
        # Give the delivery card a beat to finish rendering.
        page.wait_for_timeout(500)
        # Screenshot the smallest .bubble-element.Group that contains all of
        # 'Delivery Status' + 'Recipient' + 'Sender'. Bubble.io wraps every
        # visual container in .bubble-element.Group, so picking the smallest
        # matching one gives us the exact middle-content column — reliably,
        # unlike the old text-node walker that failed when Bubble injected
        # spans around the label. Verified against live letter 2026-07-27.
        cert_path.parent.mkdir(parents=True, exist_ok=True)
        element = None
        for attempt in range(3):
            try:
                panel_handle = page.evaluate_handle(
                    r"""() => {
                      const groups = Array.from(document.querySelectorAll('.bubble-element.Group'));
                      const matches = groups.filter(g => {
                        const t = g.innerText || '';
                        return t.includes('Delivery Status')
                            && t.includes('Recipient')
                            && t.includes('Sender');
                      });
                      if (!matches.length) return null;
                      matches.sort((a, b) => (a.innerText.length - b.innerText.length));
                      return matches[0];
                    }"""
                )
                element = panel_handle.as_element() if panel_handle else None
                if element:
                    break
                # No match yet — the letter shell may still be rendering.
                log(f"Docupost panel not found on attempt {attempt+1}; waiting 3s.")
                page.wait_for_timeout(3000)
            except PWError as e:
                log(f"Docupost panel handle transient error attempt {attempt+1}: {e}")
                page.wait_for_timeout(1500)
        if element:
            try:
                element.screenshot(path=str(cert_path))
            except PWError as e:
                log(f"Element screenshot failed, falling back to viewport: {e}")
                page.screenshot(path=str(cert_path))
        else:
            log("Docupost middle panel selector missed — falling back to viewport screenshot.")
            page.screenshot(path=str(cert_path))
        log(f"Saved Docupost certification: {cert_path}")
        return cert_path, tracking
    except Exception as e:
        log(f"Docupost certification capture failed: {type(e).__name__}: {e}")
        return None, None
    finally:
        try:
            context.close()
        except Exception:
            pass


# ------------------------------ Runner main --------------------------------

def resident_name_from_page(page: Page) -> str:
    return page.evaluate(
        r"""() => {
          const m = document.body.innerText.match(/Full name\s*\n?\s*([^\n]+)/);
          return m ? m[1].trim() : '';
        }"""
    )


def unit_number_from_page(page: Page) -> str:
    """Return the resident's unit identifier preserving any letter prefix/suffix
    (e.g. ``T135``, ``135``, ``135B``). Old form used ``/\\d+/`` which dropped
    letters — T135 came out as 135 and cascaded into wrong filenames, ResMan
    Documents naming, GitHub-hosted PDF path, and Docupost letter description."""
    return page.evaluate(
        r"""() => {
          const cells = Array.from(document.querySelectorAll('td'));
          const cell = cells.find(c => /^Unit\s+[A-Za-z]*\d+[A-Za-z]*\b/.test(c.textContent.trim()));
          return cell ? (cell.textContent.match(/^Unit\s+([A-Za-z]*\d+[A-Za-z]*)/) || ['',''])[1] : '';
        }"""
    )


def resident_email_from_page(page: Page) -> str:
    return page.evaluate(
        r"""() => {
          const a = document.querySelector('a[href^="mailto:"]');
          return a ? a.getAttribute('href').replace(/^mailto:/, '') : '';
        }"""
    )


def _find_soffice() -> str | None:
    """Locate LibreOffice's headless entrypoint on this machine."""
    for cand in ("soffice", "libreoffice", "soffice.bin"):
        p = shutil.which(cand)
        if p:
            return p
    # Windows default install path
    for p in (
        r"C:\Program Files\LibreOffice\program\soffice.exe",
        r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
    ):
        if os.path.exists(p):
            return p
    return None


def _docx_to_pdf(docx: Path) -> Path | None:
    """Convert a .docx to .pdf (same directory). Tries LibreOffice first
    (cross-platform, works on Linux CI), then docx2pdf (needs MS Word on
    Windows / MacOS). Returns the PDF path or None."""
    out_pdf = docx.with_suffix(".pdf")

    soffice = _find_soffice()
    if soffice:
        import subprocess
        log(f"docx→PDF via LibreOffice ({soffice})")
        try:
            subprocess.run(
                [soffice, "--headless", "--convert-to", "pdf",
                 "--outdir", str(docx.parent), str(docx)],
                check=True, capture_output=True, timeout=120,
            )
            if out_pdf.exists():
                return out_pdf
            log("LibreOffice ran but no PDF produced.")
        except Exception as e:
            log(f"LibreOffice conversion failed: {e}")

    try:
        from docx2pdf import convert  # type: ignore
        log("docx→PDF via docx2pdf (MS Word)")
        convert(str(docx))
        if out_pdf.exists():
            return out_pdf
    except ImportError:
        pass
    except Exception as e:
        log(f"docx2pdf conversion failed: {e}")

    return None


def merge_claim_and_fas_to_pdf(claim_docx: Path, fas_pdf: Path, out_dir: Path, unit: str, resident_name: str = "") -> Path | None:
    """Convert claim docx -> PDF then merge with FAS PDF. Skips gracefully
    if no docx→PDF converter is available on this machine.

    Output filename: `Move Out Docs - Unit <unit>.pdf` (falls back to the
    resident's slug if the unit number isn't known).
    """
    try:
        from pypdf import PdfWriter, PdfReader
    except ImportError as e:
        log(f"Skipping PDF merge (missing dependency: {e}).")
        return None

    claim_pdf = _docx_to_pdf(claim_docx)
    if not claim_pdf:
        log("Skipping PDF merge (no docx→PDF converter found — install LibreOffice or MS Word).")
        return None

    label = f"Unit {unit}" if unit else (safe_slug(resident_name) or "Unknown")
    combined = out_dir / f"Move Out Docs - {label}.pdf"

    def _is_empty_images_page(page_) -> bool:
        """ResMan's FAS PDF appends a page whose only content is the header
        text 'Images for Charges' when no charge images were uploaded. Drop
        those — but only when that's the *only* thing on the page (don't
        strip a legit page that starts with that header + actual images)."""
        txt = (page_.extract_text() or "").strip().lower()
        return txt == "images for charges"

    w = PdfWriter()
    for src in [claim_pdf, fas_pdf]:
        for page_ in PdfReader(str(src)).pages:
            if src is fas_pdf and _is_empty_images_page(page_):
                log("Skipping empty 'Images for Charges' page in FAS PDF.")
                continue
            w.add_page(page_)
    with open(combined, "wb") as f:
        w.write(f)
    try:
        claim_pdf.unlink()  # keep only the merged PDF
    except FileNotFoundError:
        pass
    log(f"Wrote merged PDF: {combined}")
    return combined


def run(payload: dict, send: bool, headless: bool) -> dict:
    lease_url = payload["leaseUrl"]
    charges   = payload["charges"]
    mor_date  = payload.get("morDate") or today_str()
    email_cfg = payload.get("email") or {}
    email_enabled = email_cfg.get("enabled", True) and send
    from_pref = email_cfg.get("from", "property")
    template  = email_cfg.get("template", DEFAULT_TEMPLATE)
    out_dir   = Path(payload.get("outputDir") or HERE / "out")
    out_dir.mkdir(parents=True, exist_ok=True)

    result: dict = {
        "status": None,
        "startedAt": now_iso(),
        "endedAt": None,
        "durationSeconds": None,
        "resident": {
            "name": None, "unit": None, "property": None,
            "leaseUrl": lease_url, "email": None,
        },
        "mor": {
            "date": mor_date, "status": None,
            "charges": [
                {"category": c.get("category", DEFAULT_CATEGORY),
                 "description": c["description"],
                 "amount": float(c["amount"])} for c in charges
            ],
            "totals": None,
            "forwardingAddress": None,
            "forwardingSource": None,
        },
        "docs": {"claimForm": None, "fasPdf": None, "combinedPdf": None},
        "email": {
            "attempted": email_enabled, "sent": False, "to": None,
            "from": from_pref, "template": template,
            "subject": None, "attachedByResMan": None,
            "commLogVerified": None, "commLogRow": None,
        },
        "docupost": None,
        "github": {
            "repo": os.environ.get("GITHUB_REPOSITORY") or "ymi-flowing/mo-process",
            "runUrl": (
                f"https://github.com/{os.environ['GITHUB_REPOSITORY']}/actions/runs/{os.environ['GITHUB_RUN_ID']}"
                if os.environ.get("GITHUB_RUN_ID") and os.environ.get("GITHUB_REPOSITORY") else None
            ),
        },
        "logs": None,      # filled at the end
        "error": None,
    }
    started = time.time()

    with sync_playwright() as pw:
        browser = pw.chromium.launch(headless=headless, args=["--start-maximized"])
        # Headed: no_viewport lets Chrome fill the actual window.
        # Headless: use an explicit large viewport so ResMan's sticky footer
        # and BalancesCell don't cover the mailto link near the bottom.
        if headless:
            context = browser.new_context(viewport={"width": 1600, "height": 1200})
        else:
            context = browser.new_context(no_viewport=True)
        page = context.new_page()

        login(page)
        mor_info = open_move_out_rec(page, lease_url, known_property_names=list(PROPERTIES.keys()))

        # Resolve which property this resident belongs to (drives Claim Form
        # return address + email, Docupost sender, and result.resident.property).
        prop_name, prop_cfg, prop_signals = resolve_property(payload, mor_info["propertySignals"])
        log(f"Resolved property: {prop_name!r}")
        result["resident"]["property"] = prop_name
        result["property"] = {
            "name":     prop_name,
            "proid":    prop_signals.get("proid"),
            "config":   prop_cfg,
        }

        fill_mor_date(page, mor_date)
        for c in charges:
            add_charge(page,
                       description=c["description"],
                       amount=float(c["amount"]),
                       category=c.get("category", DEFAULT_CATEGORY))
            page.wait_for_timeout(400)

        result["mor"]["totals"] = capture_mor_totals(page)
        log(f"Totals: {result['mor']['totals']}")

        approve_mor(page)
        result["mor"]["status"] = "Complete"

        resident_name = resident_name_from_page(page)
        result["resident"]["name"]  = resident_name
        result["resident"]["unit"]  = unit_number_from_page(page)
        result["resident"]["email"] = resident_email_from_page(page)
        log(f"Resident: {resident_name!r} unit {result['resident']['unit']!r}")

        fwd = get_forwarding_address(page)
        if fwd and fwd.get("street"):
            result["mor"]["forwardingSource"] = "resident"
        else:
            log("Forwarding blank -> falling back to unit address.")
            fwd = get_unit_address_via_new_tab(context, page) or {}
            result["mor"]["forwardingSource"] = "unit"
        result["mor"]["forwardingAddress"] = fwd

        claim_form_path = generate_claim_form(
            out_dir=out_dir,
            resident_name=resident_name,
            date_str=datetime.now().strftime("%m/%d/%Y"),
            forwarding=fwd,
            totals=result["mor"]["totals"],
            charges=charges,
            property_config=prop_cfg,
        )
        result["docs"]["claimForm"] = str(claim_form_path)

        fas_path = download_fas_pdf(page, out_dir, resident_name, mor_date)
        result["docs"]["fasPdf"] = str(fas_path) if fas_path else None

        # Merge Claim Form + FAS into one PDF *before* uploading. We upload
        # only the merged PDF to ResMan's Documents (the docx becomes an
        # intermediate) so email attachments are a single, tidy file that
        # matches what Docupost mails. Named `Move Out Docs - Unit <#>.pdf`
        # so residents' files are indexed by unit, not by their name.
        combined = None
        if fas_path:
            combined = merge_claim_and_fas_to_pdf(
                claim_form_path, fas_path, out_dir,
                unit=result["resident"]["unit"] or "",
                resident_name=resident_name,
            )
            if combined:
                result["docs"]["combinedPdf"] = str(combined)

        # ============================================================
        # NEW FLOW ORDER (2026-08-04):
        #   1. Docupost sendletter (needs merged PDF only, NOT gated on email)
        #   2. Cert screenshot from Docupost dashboard
        #   3. API-upload merged PDF to /Move-Out Docs   (HARD-FAIL)
        #   4. API-upload cert PNG   to /Move-Out Docs   (soft-fail)
        #   5. Resident email: open dialog, attach merged PDF from ResMan
        #      picker (with folder-expand + Add-from-Computer fallback), send
        #   6. Comm Log verify
        # No more Playwright button.add-files uploads — the API handles both
        # files, avoiding the accordion-visibility flakiness that broke runs
        # 30291119947 and 30858602298.
        # ============================================================

        # ------- Step 1: Docupost sendletter -------
        dp_cfg = payload.get("docupost") or {}
        dp_enabled = dp_cfg.get("enabled") is not False
        if dp_enabled and combined:
            try:
                result["docupost"] = _maybe_send_docupost(
                    cfg=dp_cfg,
                    combined_pdf=combined,
                    resident_name=result["resident"]["name"],
                    resident_unit=result["resident"]["unit"],
                    forwarding=result["mor"]["forwardingAddress"] or {},
                    repo=result["github"]["repo"],
                    property_config=prop_cfg,
                    property_name=prop_name,
                )
            except Exception as e:
                log(f"Docupost step failed: {type(e).__name__}: {e}")
                result["docupost"] = {"skipped": f"{type(e).__name__}: {e}"}
        elif dp_enabled and not combined:
            log("Docupost skipped: no Combined PDF was produced (docx→PDF failed).")
            result["docupost"] = {"skipped": "no_combined_pdf"}
        else:
            log("Docupost skipped: docupost.enabled=false in payload.")
            result["docupost"] = {"skipped": "disabled_in_payload"}

        # ------- Step 2: Mail certification screenshot from Docupost dashboard -------
        cert_path_actual: Path | None = None
        cert_tracking: str | None = None
        cert_capture_error: str | None = None
        letter_id = (result["docupost"] or {}).get("letterId")
        if letter_id:
            web_user, web_pass = _load_docupost_web_creds()
            if not (web_user and web_pass):
                cert_capture_error = "no_docupost_web_creds"
                log("Docupost cert skipped: DOCUPOST_WEB_USER/PASS not set (env or Cardentials.txt).")
            else:
                try:
                    cert_path_actual, cert_tracking = capture_docupost_certification(
                        browser=browser,
                        letter_id=letter_id,
                        out_dir=out_dir,
                        web_user=web_user,
                        web_pass=web_pass,
                    )
                    if not (cert_path_actual and cert_path_actual.exists()):
                        cert_capture_error = "screenshot_failed"
                except Exception as e:
                    cert_capture_error = f"{type(e).__name__}: {e}"
                    log(f"Docupost cert capture crashed: {e}")

        # ------- Step 3: API-upload merged PDF (HARD-FAIL) -------
        # This must succeed for the email step (attach picker looks for the
        # file in ResMan Documents). Bubble any error up to the top-level
        # exception handler which marks the run as error.
        oid = mor_info.get("oid")
        property_id = (
            (prop_cfg or {}).get("proid")
            or result["property"]["proid"]
            or prop_signals.get("proid")
        )
        merged_doc_id: str | None = None
        if combined:
            if not oid or not property_id:
                raise RuntimeError(
                    f"Cannot API-upload merged PDF: missing oid={oid!r} or propertyId={property_id!r}"
                )
            resp = upload_document_via_api(combined, object_id=oid, property_id=property_id)
            merged_doc_id = resp.get("documentId")
            result["docs"]["combinedPdfDocumentId"] = merged_doc_id

        # ------- Step 4: API-upload cert PNG (SOFT-FAIL) -------
        cert_uploaded = False
        cert_upload_error = cert_capture_error
        cert_doc_id: str | None = None
        if cert_path_actual and cert_path_actual.exists() and oid and property_id:
            try:
                resp = upload_document_via_api(
                    cert_path_actual, object_id=oid, property_id=property_id,
                )
                cert_uploaded = True
                cert_doc_id = resp.get("documentId")
            except Exception as e:
                cert_upload_error = f"upload_failed: {type(e).__name__}: {e}"
                log(f"Cert PNG API upload failed (continuing): {e}")

        result["docupost"] = {
            **(result["docupost"] or {}),
            "certification": {
                "uploaded":   cert_uploaded,
                "path":       str(cert_path_actual) if cert_uploaded and cert_path_actual else None,
                "documentId": cert_doc_id,
                "tracking":   cert_tracking,
                "error":      cert_upload_error,
            },
        }

        # ------- Step 5: Resident email — attach merged PDF, send -------
        if email_enabled:
            # ResMan indexes API-uploaded docs quickly, but give a small
            # settle so the picker inventory is warm.
            page.wait_for_timeout(2000)
            open_send_email_dialog(page)
            set_from(page, from_pref)
            apply_template(page, template)
            set_from(page, from_pref)  # template resets From; re-set.
            if combined:
                attachment_names = [combined.name]
                fallback_paths = [combined]
            else:
                attachment_names = [claim_form_path.name]
                fallback_paths = [claim_form_path]
                if fas_path:
                    attachment_names.append(fas_path.name)
                    fallback_paths.append(fas_path)
            result["email"]["attachedByResMan"] = attach_from_resman(
                page, attachment_names, fallback_local_paths=fallback_paths,
            )
            # Attaching often re-triggers the From default; re-set.
            set_from(page, from_pref)
            result["email"]["to"] = result["resident"]["email"]
            result["email"]["subject"] = f"{result['resident']['property']} - Move-Out Documents"
            click_send(page)
            result["email"]["sent"] = True

            # ------- Step 6: Verify via Communication Log -------
            subject_hint = result["email"]["subject"] or prop_name
            comm = verify_via_comm_log(page, lease_url, subject_hint, wait_seconds=10)
            if not comm["verified"]:
                log("Comm Log verify miss; retrying with longer wait.")
                comm = verify_via_comm_log(page, lease_url, subject_hint, wait_seconds=15)
            result["email"]["commLogVerified"] = comm["verified"]
            result["email"]["commLogRow"] = comm["row"]
        else:
            log("Email skipped (either --no-send or email.enabled=false).")

        context.close()
        browser.close()

    if not result["email"]["sent"]:
        result["status"] = "parked" if not email_enabled else "sent_no_email"
    else:
        # commLogVerified is None when verification wasn't run (email disabled)
        # or True/False when it did run. Treat False as "we clicked Send but
        # cannot confirm the message left ResMan" — same category as
        # sent_no_email so downstream (n8n summary email) flags it.
        result["status"] = "sent" if result["email"]["commLogVerified"] is not False else "sent_no_email"
    result["endedAt"] = now_iso()
    result["durationSeconds"] = int(time.time() - started)
    return result


def _maybe_send_docupost(
    cfg: dict,
    combined_pdf: Path,
    resident_name: str,
    resident_unit: str,
    forwarding: dict,
    repo: str,
    property_config: dict | None = None,
    property_name: str | None = None,
) -> dict:
    """Push the Combined PDF to the public repo, then hand its raw URL to
    Docupost's sendletter API. Returns the docupost result block or a
    {'skipped': reason} dict."""
    token_docupost = os.environ.get("DOCUPOST_TOKEN")
    token_gh       = os.environ.get("GITHUB_TOKEN")

    if not token_docupost:
        log("Docupost skipped: DOCUPOST_TOKEN env var missing.")
        return {"skipped": "no_docupost_token"}
    if not token_gh:
        log("Docupost skipped: GITHUB_TOKEN missing (only available in Actions).")
        return {"skipped": "no_github_token"}

    # Recipient must have street + city + state + zip.
    required = ("street", "city", "state", "zip")
    if not all(forwarding.get(k) for k in required):
        log(f"Docupost skipped: forwarding address incomplete ({[k for k in required if not forwarding.get(k)]})")
        return {"skipped": "incomplete_address"}

    # 1. Push the PDF and get its public raw URL.
    raw_url = push_pdf_to_repo(combined_pdf, repo=repo, token=token_gh)
    if not _wait_raw_url_live(raw_url):
        return {"skipped": "raw_url_not_live", "pdfUrl": raw_url}

    # 2. Build sender + recipient dicts. Property-directory config wins over
    # DEFAULT_SENDER; explicit cfg.sender still wins over both.
    prop_sender = {}
    if property_config:
        prop_sender = {
            "name":     property_name or DEFAULT_SENDER["name"],
            "address1": property_config.get("address1"),
            "city":     property_config.get("city"),
            "state":    property_config.get("state"),
            "zip":      property_config.get("zip"),
        }
        prop_sender = {k: v for k, v in prop_sender.items() if v}
    sender = { **DEFAULT_SENDER, **prop_sender, **(cfg.get("sender") or {}) }

    # Split "8400 49th Street North Apt. 1113" into street1 + optional apt.
    street = forwarding["street"]
    address1, address2 = street, None
    m = re.match(r"^(.*?)\s+(Apt\.?\s*\S+|Unit\s*\S+|#\s*\S+)$", street, re.I)
    if m:
        address1, address2 = m.group(1).strip(), m.group(2).strip()

    recipient = {
        "name":     resident_name or "",
        "address1": address1,
        "address2": address2,
        "city":     forwarding["city"],
        "state":    forwarding["state"],
        "zip":      forwarding["zip"],
    }

    cfg = dict(cfg)  # avoid mutating payload
    cfg.setdefault("description", f"MO {resident_unit or ''} {resident_name or ''}".strip()[:40])

    return send_via_docupost(cfg, sender, recipient, raw_url, token_docupost)


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--payload", required=True, help="JSON, '@file', or '-' for stdin")
    ap.add_argument("--headless", action="store_true", help="Run headless (default: headed)")
    ap.add_argument("--no-send", action="store_true", help="Skip the final email Send click")
    args = ap.parse_args()

    started_iso = now_iso()
    started_t   = time.time()

    payload = None
    try:
        payload = load_payload(args.payload)
        result = run(payload, send=not args.no_send, headless=args.headless)
        result["logs"] = list(_LOGS)
    except Exception as e:
        log(f"FATAL: {type(e).__name__}: {e}")
        result = {
            "status": "error",
            "startedAt": started_iso,
            "endedAt":   now_iso(),
            "durationSeconds": int(time.time() - started_t),
            "resident": {
                "name": None, "unit": None, "property": None,
                "leaseUrl": (payload or {}).get("leaseUrl") if isinstance(payload, dict) else None,
                "email": None,
            },
            "mor":   {"date": None, "status": None, "charges": [], "totals": None,
                      "forwardingAddress": None, "forwardingSource": None},
            "docs":  {"claimForm": None, "fasPdf": None, "combinedPdf": None},
            "email": {"attempted": False, "sent": False, "to": None,
                      "from": None, "template": None, "subject": None,
                      "attachedByResMan": None,
                      "commLogVerified": None, "commLogRow": None},
            "docupost": None,
            "github": {
                "repo":   os.environ.get("GITHUB_REPOSITORY") or "ymi-flowing/mo-process",
                "runUrl": (
                    f"https://github.com/{os.environ['GITHUB_REPOSITORY']}/actions/runs/{os.environ['GITHUB_RUN_ID']}"
                    if os.environ.get("GITHUB_RUN_ID") and os.environ.get("GITHUB_REPOSITORY") else None
                ),
            },
            "logs":  list(_LOGS),
            "error": {
                "type": type(e).__name__,
                "message": str(e),
                "traceback": traceback.format_exc().splitlines()[-10:],
            },
        }

    # Emit a single JSON result on stdout so callers/CI can capture it.
    print(json.dumps(result, indent=2, default=str))
    # Non-zero exit on error so GH Actions marks the job failed but still emits JSON.
    sys.exit(1 if result.get("status") == "error" else 0)


if __name__ == "__main__":
    main()
